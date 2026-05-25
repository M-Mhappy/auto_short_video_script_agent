import os
import uuid
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from backend.tools.output_paths import book_output_dirs, book_slug


def export_word_node(state: dict) -> dict:
    script = state.get("script_draft", "")
    chapters = state.get("script_chapters") or []
    book = state.get("selected_book", {})
    title = book.get("title", "口播稿")
    author = book.get("author", "")

    dirs = book_output_dirs(title)
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "微软雅黑"
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    heading = doc.add_heading(f"《{title}》口播稿", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if author:
        sub = doc.add_paragraph(f"原著：{author}")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].font.size = Pt(10)

    doc.add_paragraph("")

    if chapters:
        for ch in chapters:
            ch_title = ch.get("title", "").strip()
            ch_content = ch.get("content", "").strip()
            if ch_title:
                doc.add_heading(ch_title, level=2)
            for para_text in ch_content.split("\n"):
                stripped = para_text.strip()
                if stripped:
                    doc.add_paragraph(stripped)
    else:
        for para_text in script.split("\n"):
            stripped = para_text.strip()
            if stripped:
                if stripped.startswith("## "):
                    doc.add_heading(stripped[3:].strip(), level=2)
                else:
                    doc.add_paragraph(stripped)

    prefix = book_slug(title)
    filename = f"{prefix}_口播稿_{uuid.uuid4().hex[:8]}.docx"
    filepath = os.path.join(dirs["word"], filename)
    doc.save(filepath)

    return {
        "word_file_path": filepath,
        "current_step": "export_word",
    }
